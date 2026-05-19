"""ブログ記事をWord文書として作成（新構成版）"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- スタイル設定 ---
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5

for level, size, color in [(1, 18, '1a1a2e'), (2, 15, '16213e'), (3, 13, '0f3460')]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Yu Gothic'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor.from_string(color)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(24 if level == 1 else 18)
    h.paragraph_format.space_after = Pt(12)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
for attr in ['left_margin', 'right_margin', 'top_margin', 'bottom_margin']:
    setattr(section, attr, Cm(2.5))

BLOG_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_WIDTH = Inches(5.8)

def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('1a1a2e')
    run.font.name = 'Yu Gothic'
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(20)
        r = p2.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)
        r.font.name = 'Yu Gothic'

def add_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    return p

def add_image(filename, caption=None):
    img_path = os.path.join(BLOG_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=IMG_WIDTH)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(16)
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)
            r.font.name = 'Yu Gothic'

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Yu Gothic'

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('0f3460')

def add_note(text):
    """注記（グレーのイタリック）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run('― ― ―')
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(14)


# ========================================
# 記事本文
# ========================================

add_title(
    '「"辛口がいい酒"の時代は終わったのか？」',
    '― 140年分のデータをつないで見えた、日本酒トレンドの現在地'
)

add_text('「日本酒は辛口にかぎる」――以前は居酒屋でもよく聞かれたこの言い回しが、いまは少し古く見える。ランキング上位に並ぶのは、甘口でフルーティな酒ばかりに見えるからだ。では、本当に"辛口の時代"は終わったのだろうか。')

add_text('国立国会図書館の出版物データ、日本酒レビューサイトSAKETIMEの約37万件のレビュー、そしてGoogle Trendsの検索データをつないで追ってみると、見えてきたのは単純な「辛口から甘口へ」という物語ではなかった。上位ランキングの景色と、一般消費者が酒を探すときの言葉は、必ずしも一致していない。')

add_separator()

# ===== 出発点 =====
doc.add_heading('出発点：あるYouTube動画の主張', level=1)

add_text('きっかけはYouTubeの日本酒動画だった。そこでは、SAKETIMEのランキング上位に並ぶ銘柄を眺めて、「上位の日本酒は、甘旨（あまうま）ジューシーなものばかり」という結論に達していた。')

add_text('十四代、新政、而今、花陽浴――どの銘柄も一言では語り尽くせない深みを持っているが、直感としては納得できる観察だ。しかし、それをそのまま「日本酒全体のトレンド」と呼んでよいのかは別問題である。そして、もしランキング上位が本当に甘口ばかりだとしたら、あの「辛口がいい酒」という長年の常識はどうなったのか。')

add_text('そこで本稿では、まず長期の歴史的な言葉の流れを見たうえで、次に愛好家コミュニティの評価構造を確認し、最後に一般消費者の検索行動と照らし合わせることにした。狙いは、ひとつのデータで結論を急がず、異なる層で"辛口"がどう生き残っているかを確かめることにある。')

add_separator()

# ===== 第1章 =====
doc.add_heading('1. 「辛口がいい酒」はいつから強くなったのか', level=1)

add_text('長いスパンで見るため、まず国立国会図書館の NDL Ngram Viewer で「辛口」「甘口」「淡麗」「芳醇」「吟醸」「純米」の出現比率を確認した。ここで注意したいのは、NDLのデータは日本酒に限定されておらず、出版物全体の語頻度であることだ。したがってこれは厳密な"日本酒史"ではなく、あくまで時代の言葉づかいの変化を見るための補助線と考えるのがよい。')

add_image('fig8_ndl_ngram.png', '図1: 出版物における味わい語の出現比率（1880〜2000年）。出典: NDL Ngram Viewer。')

add_text('この長期データでは、明治期（1880年代）には「甘口」が圧倒的に多く使われていた。文献ではよく指摘されることだが、江戸時代の日本酒の味わいの語りはもともと「甘口」が中心だった。そうした傾向がデータに現れている可能性がある。')

add_text('その後「甘口」は100年かけて緩やかに減少し、代わりに1970年代から「辛口」が急上昇する。同じ時期に日本酒固有の語である「淡麗」も伸びているため、少なくとも言説のレベルでは、いわゆる「淡麗辛口」ブームがこの時代に強く立ち上がったと読むのが自然だ。')

add_text('なお、「辛口」信仰を戦後の三増酒（米不足から醸造アルコール、糖類、酸味料を加えて原料米の約3倍量に増量した酒。甘くベタつく味わいから低品質な酒の代名詞とされた）へのカウンターとして説明する議論もよく聞かれる。しかし、少なくとも第二次世界大戦直後に「辛口」が文献レベルで増えたという証拠は、このデータからは見つけられなかった。')

add_text('下段のグラフでは、「吟醸」が1980年代から爆発的に増加している。吟醸酒ブームの実態がデータで裏付けられた格好だ。逆にいえば、「辛口がいい酒」という感覚は、日本酒の長い歴史全体から見れば比較的新しい価値観である。戦後直後から一貫して支配的だったというより、1970〜80年代に強く制度化・一般化した語りとみるほうが実態に近い。')

add_separator()

# ===== 第2章 =====
doc.add_heading('2. 愛好家コミュニティの上位銘柄は、たしかに甘口寄りだった', level=1)

add_text('次に、SAKETIMEの全5,386銘柄・約37万件のレビューを見ていく。ここで把握できるのは「市場全体」ではなく、レビューを書く愛好家コミュニティの選好である。それでも、いま何が高く評価されやすいかを知るには強力な観測点だ。')

doc.add_heading('構造化データで見ると、上位ほど甘口・軽やか', level=2)

add_text('SAKETIMEでは、レビュー時に「甘辛」（辛い+2〜甘い+2の5段階）と「ボディ」（軽い+2〜重い+2の5段階）を登録できる。これをランキング帯別に集計すると、上位銘柄の輪郭はかなりはっきりしていた。')

add_image('fig2_verification_structured.png', '図2: ランキング帯別の甘口率と濃醇率。')

add_bullet('甘口率（甘い+1以上）: Top10で67.3%、101-500位で30.2%、501位以降で27.6%')
add_bullet('辛口率（辛い+1以上）: Top10でわずか3.7%、101-500位で25.7%、501位以降で29.3%')
add_bullet('濃醇率（重い+1以上）: Top10で11.2%、101-500位で18.4%')

add_text('甘口率はTop10（67%）と501位以降（28%）で実に40ポイントもの差がある。逆に辛口率を見ると、501位以降（29%）はTop10（4%）の7倍以上だ。ランキング上位の世界では「甘口＝高評価」になりやすい、と言ってよい結果だ。')

add_text('ただしそれは「濃醇で旨味が重い」こととは同義ではない。濃醇率はむしろ上位ほど低く、上位に集まりやすいのは、甘さがありつつも重すぎず、輪郭がきれいで飲み疲れしないタイプだ。YouTube動画の「甘口」という指摘は正しかったが、「旨味」「ジューシー」については慎重に扱う必要がある。')

doc.add_heading('レビューテキストには「フルーティ」の強さが現れる', level=2)

add_text('構造化データだけでは見えない言葉のニュアンスを確かめるため、レビューテキストのキーワード出現率も比較した。')

add_image('fig3_verification_text.png', '図3: レビューテキスト中の味わいキーワード出現率（ランキング帯別）')

add_bullet('甘口系キーワード: Top10で32.8%、501位以降で26.8%')
add_bullet('辛口系キーワード: Top10で3.7%、501位以降で16.8%（上位ほど圧倒的に少ない）')
add_bullet('フルーティ系: Top10で24.7%、501位以降で13.3%（上位で約2倍）')
add_bullet('芳醇系: Top10で17.9%、501位以降で23.2%（むしろ上位が低い）')

add_text('ここで目立つのは、上位ほど「辛口」が語られにくく、「フルーティ」が語られやすいことだ。甘口系キーワードも上位でやや多いが、それ以上に、上位銘柄の印象を特徴づけているのは"果実感のある華やかさ"だと言ってよい。')

add_image('fig4_verification_radar.png', '図4: Top10銘柄の味わいプロファイル。')

add_text('Top10の個別プロファイルを見ると、花陽浴のフルーティ率37%、陽乃鳥・宮寒梅の甘口率44%が際立つ。一方で、ランキング1位の十四代は辛口率わずか1.8%で、甘口率も特別高いわけではない、極端な偏りの少ないバランス型として立っている。つまり、上位銘柄の美学は「甘ければよい」ではなく、「華やかさと上品さをどう両立させるか」に近い。')

add_image('fig5_verification_heatmap.png', '図5: テイスト分布の比較。')

add_text('甘辛とボディの分布を重ねると、上位銘柄の中心地は「普通ボディ × 甘い寄り」に集まる。ここから見えてくるのは、"重厚芳醇"の勝利ではなく、"ほどよく軽く、香りが立ち、甘さがきれいに見える酒"の優位である。')

add_separator()

# ===== 第3章 =====
doc.add_heading('3. ただし、甘口一辺倒とまでは言い切れない', level=1)

add_text('では、コミュニティ全体もひたすら甘口化しているのか。2017〜2025年の時系列で見ると、景色はもう少し複雑だ。')

doc.add_heading('構造化データでは、甘口は定着しつつも独走ではない', level=2)

add_image('fig6_trend_structured.png', '図6: テイスト傾向の時系列推移（構造化データ、全銘柄）')

add_bullet('甘口率は期間を通じて最も高い水準を維持しているが、2017年の42%から2025年の37%へ緩やかに低下している')
add_bullet('辛口率は2022年ごろまで下がったあと、直近でやや持ち直している（18% → 20%台へ反転）')
add_bullet('濃醇率は20%→14%へ下がり続けており、「重い酒」離れの流れは比較的はっきりしている')
add_bullet('軽快率は28-30%で安定')

add_text('ここで重要なのは、「甘口が主流になった」という観察と、「辛口が終わった」という結論は同じではないという点だ。少なくともコミュニティ全体では、甘口が王道化した一方で、辛口も一定の存在感を保っている。')

doc.add_heading('テキスト分析でも、伸びているのは甘口以上に"フルーティ"', level=2)

add_image('fig7_trend_text.png', '図7: 味わいキーワード出現率の推移（テキスト分析、全銘柄）')

add_text('レビューテキストでも、最もはっきり伸びているのはフルーティ系の表現だった。2017年の17%から2025年には25%まで上昇している。甘口系はすでに主流として定着したため伸びが鈍くなり、辛口系は一時的な低下のあとに下げ止まりが見える。')

add_quote('現状を表すなら「辛口から甘口へ完全移行した」というより、「甘口・フルーティが上位の標準語になった一方で、コミュニティ全体では多様化が進んでいる」と言うほうが正確だろう。')

add_separator()

# ===== 第4章 =====
doc.add_heading('4. 一般消費者の検索軸では、いまも「辛口」が強い', level=1)

add_text('最後にGoogle Trendsを見る。ここで測っているのは味覚そのものではなく、"人が酒を探すときに使う言葉"である。したがって、これは選好の直接測定というより、分類語・検索語としての強さを見るデータだ。')

add_image('fig9_google_trends.png', '図8: Google Trendsにおける日本酒関連の検索トレンド（2004〜2025年）')

add_text('このデータでは、「日本酒 辛口」が「日本酒 甘口」を一貫して大きく上回る。しかも、辛口は2004年から2025年にかけて長期的に右肩上がりで、甘口も緩やかに増えている。つまり一般消費者の検索行動では、"辛口"はいまも非常に強い入口語であり続けている。')

add_text('もちろん、検索量がそのまま好みの強さを意味するわけではない。辛口は店頭でも通じやすく、無難な選び方のラベルとして機能している可能性が高い。それでも、少なくとも「辛口」という言葉がすでに死語である、とはとても言えない。')

doc.add_heading('三つの層', level=2)

add_text('ここまでの観察を整理すると、日本酒の味わいをめぐる風景は三層構造になっている。')

add_bullet('第1層・ランキング上位（Top10〜50位）：甘口・フルーティ・軽快が支配的。甘口率67%、辛口率4%。')
add_bullet('第2層・ランキング下位（501位以降）：甘口率28%、辛口率29%。甘辛が拮抗しており、辛口の銘柄も数多く存在。')
add_bullet('第3層・一般消費者（Google検索）：「辛口」が一貫して主要キーワード。検索量は増加中。')

add_text('しかも、前章で見たように、SAKETIMEのレビュー全体でも辛口率は近年再上昇している。愛好家コミュニティの中でさえ、辛口への揺り戻しが始まっているのだ。')

add_separator()

# ===== 結論 =====
doc.add_heading('5. 結論：「辛口の時代」は終わったのか', level=1)

add_text('上位ランキングだけを見れば、答えはたしかに「かなり終わったように見える」だ。いま高く評価されやすいのは、甘口で、フルーティで、重すぎない酒だからである。')

add_text('しかし視野を広げると、答えは変わる。ランキングの裾野には辛口の銘柄がなお多く存在し（501位以降では辛口率29%）、Google検索では辛口が依然として主力の探索語であり（甘口の2.5倍）、コミュニティ全体でも辛口は完全には後退しきっていない。見えてくるのは、「甘口が王道になった」のではなく、「上位銘柄の評価軸が変わった」ということ、そしてその変化が市場全体の言葉づかいと完全には一致していないということだ。')

add_text('約140年のスパンで俯瞰すると、甘口（明治）→ 辛口（1970-90年代）→ 甘口・フルーティ（2010年代後半〜）という大きな揺り戻しのサイクルが見える。ただし現在起きていることは単純な「甘口回帰」ではない。甘口・フルーティが王道として定着する一方で、辛口も復権の兆しを見せており、むしろ味わいの多様化が進んでいるようにも見える。')

add_quote('"辛口がいい酒"の時代は、少なくとも終わってはいない。ただしその言葉が、もはや上位銘柄の価値を最もよく説明する言葉でもなくなっている。いまの日本酒は、甘口・フルーティを中心に再編されつつ、それでも辛口を捨てきってはいない。')

add_separator()

# ===== 補論 =====
doc.add_heading('補論1：SAKETIME全体像と、ランキング解釈の注意点', level=1)

add_text('SAKETIMEの分布を俯瞰すると、評価スコアとレビュー数の相関が非常に高い（対数スケールで0.888）。これは単に「おいしい酒ほどレビューされる」だけでなく、「知られた酒ほど高く評価されやすい」「話題性のある酒が上位に残りやすい」といった可視性バイアスを含んでいる可能性がある。上位ランキングを"市場全体の縮図"とみなしてよいわけではない。')

add_image('fig1_rating_vs_reviews.png', '補図1: 評価スコア vs レビュー数（上位10銘柄をラベル表示）')

add_text('この点は、本稿の中心結論を読むうえで重要だ。ここで見ているのは「いま評価が集まりやすい銘柄の味わい」であり、日本酒市場全体の販売実績や一般消費者の飲用実態を直接測っているわけではない。')

add_separator()

doc.add_heading('補論2：特定名称による評価の違い', level=1)

add_text('特定名称別に見ると、純米大吟醸（4.16）や純米吟醸（4.05）の平均評価が高く、本醸造（3.79）や普通酒（3.75）は相対的に低い。これ自体は直感に沿う結果だが、差は0.4ポイント程度と決定的というほど大きくはなく、結局のところ名称だけで評価が決まっているわけではない。')

add_image('fig10_price_specname.png', '補図2: 特定名称別の平均評価')

add_separator()

doc.add_heading('補論3：さけのわデータによる独立検証', level=1)

add_text('本文の結論がSAKETIME特有の偏りに引っ張られていないかを確認するため、公開APIが利用できる「さけのわ」のフレーバーデータでもクロスチェックした。さけのわは日本酒コミュニティアプリで、100万件以上のチェックイン記録を持つ。約3,100銘柄のフレーバーチャート（華やか/芳醇/重厚/穏やか/軽快/ドライの6軸数値）とフレーバータグ（141種類）が公開されている。')

add_image('fig12_sakenowa_analysis.png', '補図3: さけのわランキング帯別のフレーバータグ出現率（左）とフレーバーチャート6軸平均（右）')

add_text('結果はおおむね一致していた。フレーバーチャート6軸を見ると、上位銘柄ほど「華やか」が高く（1-10位: 0.50 → 圏外: 0.35）、「重厚」が低い（1-10位: 0.24 → 圏外: 0.35）。フレーバータグでも、辛口タグが1-10位で0%なのに対し51-100位では78%に達する。フルーティは1-10位で100%、圏外で33%。')

add_text('辛口や重さが前面に出るのはむしろ中位以下であり、上位では"華やかで軽やか、しかもバランスがよい"という方向が再現される。少なくとも「上位ほどフルーティ寄りで、重厚辛口ではない」という骨格は、別データでも支持された。')

add_separator()

# ===== 分析手法とデータソース =====
doc.add_heading('分析手法とデータソースについて', level=1)
add_bullet('SAKETIMEデータ: Pythonによるスクレイピング（BeautifulSoup）。5,386銘柄のランキング情報、銘柄詳細、約37万件のレビューを取得')
add_bullet('テキスト分析: キーワード出現率の集計、ランキング帯別・年代別クロス分析')
add_bullet('さけのわデータ: さけのわデータプロジェクトAPI（約3,100銘柄のフレーバーチャート6軸・フレーバータグ141種）。さけのわデータ (https://sakenowa.com) のデータを加工して利用')
add_bullet('出版物データ: 国立国会図書館 NDL Ngram Viewer API（図書・雑誌約230万点のOCRテキスト）')
add_bullet('検索トレンド: Google Trends（pytrends、地域:日本、2004年〜2025年）')

add_note('【NDL Ngram Viewerについて】NDL Ngram Viewerのデータは全出版物における単語の出現頻度であり、日本酒の文脈に限定されていない。たとえば「辛口」には批評の「辛口コメント」、「甘口」にはカレーの「甘口」なども含まれる。日本酒固有の語である「淡麗」「吟醸」「純米」は比較的日本酒文脈に近いと考えられるが、厳密には他の用法も含まれる点に留意されたい。あくまで長期的なトレンドの参考値として解釈すべきデータである。')

add_note('【さけのわフレーバータグについて】さけのわのフレーバータグはSAKETIMEのテイスト評価（甘い/辛いの排他的選択）とは性質が異なる。レビューテキストから抽出された特徴語の集積であり、1銘柄に平均9.4個のタグが付く。「甘味」と「辛口」が同時に付いている銘柄も30%存在する。タグの出現率はその味わいが言及される頻度を示すものであり、排他的な分類ではない。')

add_note('【レビュアーの偏りについて】SAKETIMEの約37万件のレビューは4,625人のユーザーによって投稿されているが、その分布は極めて偏っている。最多投稿者は1人で約14,000件（全体の3.8%）を投稿しており、上位100人（全ユーザーの2.2%）が全レビューの37%を、上位500人（10.8%）が72%を占める。本分析の結果は、こうした一部のヘビーレビュアーの嗜好に強く影響されている可能性がある点に留意されたい。')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('データ取得日: 2026年4月14日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(130, 130, 130)
run.font.name = 'Yu Gothic'

# 保存
output_path = os.path.join(BLOG_DIR, 'article_v2.docx')
doc.save(output_path)
print(f'保存完了: {output_path}')
