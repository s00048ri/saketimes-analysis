"""ブログ記事をWord文書として作成（新構成版）"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- スタイル設定 ---
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5

for level, size, color in [(1, 18, '1a1a2e'), (2, 15, '16213e'), (3, 13, '0f3460')]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Yu Gothic'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor.from_string(color)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(24 if level == 1 else 18)
    h.paragraph_format.space_after = Pt(12)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
for attr in ['left_margin', 'right_margin', 'top_margin', 'bottom_margin']:
    setattr(section, attr, Cm(2.5))

BLOG_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_WIDTH = Inches(5.8)

def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('1a1a2e')
    run.font.name = 'Yu Gothic'
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(20)
        r = p2.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)
        r.font.name = 'Yu Gothic'

def add_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    return p

def add_image(filename, caption=None):
    img_path = os.path.join(BLOG_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=IMG_WIDTH)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(16)
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)
            r.font.name = 'Yu Gothic'

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Yu Gothic'

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string('0f3460')

def add_note(text):
    """注記（グレーのイタリック）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run('― ― ―')
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(14)


# ========================================
# 記事本文
# ========================================

add_title(
    '「"辛口がいい酒"の時代は終わったか？」',
    '― 140年分のデータが示す意外な答え'
)

add_text('「日本酒は辛口がうまい」―― 居酒屋でもよく聞くこのフレーズ。しかし最近、日本酒レビューサイトのランキング上位を占めるのは甘口でフルーティな銘柄ばかりだ。辛口の時代は本当に終わったのか？ 国立国会図書館の出版物データ、日本酒レビューサイトSAKETIMEの約37万件のレビュー、そしてGoogle Trendsの検索データを横断分析し、140年のスパンで答えを探った。')

add_separator()

# ===== きっかけ =====
doc.add_heading('きっかけ：あるYouTube動画の主張', level=1)

add_text('きっかけはYouTubeの日本酒動画だった。その動画では、日本最大級の日本酒レビューサイト「SAKETIME」のランキングを見て、「上位に入る日本酒は、甘口で芳醇なものが多い」という結論に達していた。')

add_text('直感的にはそうかもしれない。しかし、本当にデータで裏付けられるのだろうか？ そして、もしランキング上位が甘口ばかりだとしたら、あの「辛口がいい酒」という長年の常識はどうなったのか？')

add_text('そこで、SAKETIMEの全5,386銘柄・約37万件のレビューをスクレイピングし、体系的に検証することにした。さらに、「辛口信仰」がそもそもいつ始まったのかを、国立国会図書館の出版物データやGoogle Trendsで140年のスパンで追ってみた。')

add_separator()

# ===== 第1章 =====
doc.add_heading('第1章：「辛口がいい酒」はいつ始まったのか', level=1)

add_text('まず、長い歴史のスパンで見てみよう。国立国会図書館が提供するNDL Ngram Viewerは、約230万点の図書・雑誌のOCRテキストから、キーワードの出現頻度を年代別に可視化できるツールだ。「辛口」「甘口」「淡麗」「芳醇」「吟醸」「純米」の出現比率を1880年代から追った。')

add_note('【注意】NDL Ngram Viewerのデータは全出版物における単語の出現頻度であり、日本酒の文脈に限定されていない。たとえば「辛口」には批評の「辛口コメント」、「甘口」にはカレーの「甘口」なども含まれる。日本酒固有の語である「淡麗」「吟醸」「純米」は比較的日本酒文脈に近いと考えられるが、厳密には他の用法も含まれる点に留意されたい。あくまで長期的なトレンドの参考値として解釈すべきデータである。')

add_image('fig8_ndl_ngram.png', '図1: 出版物における味わい語の出現比率（1880〜2000年）。出典: NDL Ngram Viewer。日本酒以外の文脈も含む。')

add_text('結果は驚くべきものだった。明治時代（1880年代）には、「甘口」が出版物で圧倒的に多く使われていた。つまり、日本酒の味わいの語りはもともと「甘口」が中心だったのだ。')

add_text('「甘口」はその後100年かけて緩やかに減少し、代わりに1970年代から「辛口」が急上昇する。同時期に日本酒固有の語である「淡麗」も上昇していることから、これは日本酒の「淡麗辛口ブーム」の到来を反映していると推測できる。')

add_text('下段のグラフでは、「吟醸」が1980年代から爆発的に増加している。吟醸酒ブームの実態がデータで裏付けられた格好だ。「辛口がいい酒」という信仰は、実はこの1970年代〜80年代に形成された、歴史的に見れば比較的新しい価値観なのである。')

add_separator()

# ===== 第2章 =====
doc.add_heading('第2章：いま、ランキング上位を占めるのはどんな酒か', level=1)

add_text('では、現在の日本酒レビューサイトではどうなっているのか。SAKETIMEの全5,386銘柄のランキングデータを分析した。')

doc.add_heading('ランキングの全体像', level=2)

add_text('評価スコアの平均は3.18（5点満点）で、大半が3.0前後に集中している。注目すべきは、レビュー数（対数スケール）と評価の相関が0.888と非常に高いこと。多くのレビューを集める銘柄ほど高評価を得ている。')

add_image('fig1_rating_vs_reviews.png', '図2: 評価スコア vs レビュー数。上位10銘柄をラベル表示。')

doc.add_heading('参考：特定名称による評価の違い', level=2)

add_image('fig10_price_specname.png', '図：特定名称別の平均評価')

add_text('参考までに、特定名称酒のカテゴリーが評価にどう影響するかを見ておこう。純米大吟醸（4.16）が最も高く、純米吟醸（4.05）が続く。一方、本醸造（3.79）や普通酒（3.75）は相対的に低い。精米歩合が高く手間をかけた酒ほど高評価という、ある意味で順当な結果だ。ただし差は0.4ポイント程度であり、特定名称だけで評価が決まるわけではない。')

add_separator()

doc.add_heading('構造化データ：上位ほど甘口、しかし「芳醇」ではない', level=2)

add_text('SAKETIMEではレビュー時に「甘辛」（辛い+2〜甘い+2の5段階）と「ボディ」（軽い+2〜重い+2の5段階）を選択できる。これをランキング帯別に集計した結果、明確な傾向が見えた。')

add_image('fig2_verification_structured.png', '図3: ランキング帯別の甘口率と濃醇率。')

add_bullet('甘口率（甘い+1以上）: Top10で67.3%、101-500位で30.2%、501位以降で27.6%')
add_bullet('辛口率（辛い+1以上）: Top10でわずか3.7%、101-500位で25.7%、501位以降で29.3%')
add_bullet('濃醇率（重い+1以上）: Top10で11.2%、101-500位で18.4% → 上位ほどむしろ軽い')

add_text('甘口率はTop10（67%）と501位以降（28%）で実に40ポイントもの差がある。逆に辛口率を見ると、501位以降（29%）はTop10（4%）の7倍以上だ。ランキングの上位と下位では、「甘口の世界」と「辛口の世界」がはっきりと分かれている。')

add_text('YouTube動画の「甘口」という指摘は正しかったが、「芳醇」は誤りだった。上位銘柄は甘口だが、ボディは重厚ではなく普通〜やや軽めなのだ。')

doc.add_heading('テキスト分析：浮かび上がる「フルーティ」の存在感', level=2)

add_text('約37万件のレビューテキストから味わいキーワードの出現率をランキング帯別に比較すると、もう一つの特徴が見えてきた。')

add_image('fig3_verification_text.png', '図4: レビューテキスト中の味わいキーワード出現率（ランキング帯別）')

add_bullet('甘口系キーワード: Top10で32.8%、501位以降で26.8%')
add_bullet('辛口系キーワード: Top10で3.7%、501位以降で16.8%（上位ほど圧倒的に少ない）')
add_bullet('フルーティ系: Top10で24.7%、501位以降で13.3%（上位で約2倍）')
add_bullet('芳醇系: Top10で17.9%、501位以降で23.2%（むしろ上位が低い）')

add_text('Top10銘柄の味わいプロファイルをレーダーチャートにすると、各銘柄の個性が浮かび上がる。')

add_image('fig4_verification_radar.png', '図5: Top10銘柄の味わいプロファイル。花陽浴のフルーティ率37%が突出。')

add_text('花陽浴のフルーティ率37%、陽乃鳥・宮寒梅の甘口率44%が際立つ。一方、ランキング1位の十四代は辛口率わずか1.8%で、突出した偏りがなくバランスの良さが特徴だ。')

doc.add_heading('テイスト分布：上位の「典型的プロファイル」', level=2)

add_text('ボディ × 甘辛のクロス集計をヒートマップにすると、上位銘柄の味わいの中心地が見える。')

add_image('fig5_verification_heatmap.png', '図6: テイスト分布の比較。Top10は「普通ボディ×甘い+1」に集中。')

add_quote('いまのランキング上位は「甘口でフルーティ、ボディは普通〜やや軽めで上品」。「重厚で芳醇」ではなく「適度な軽さの甘口」が上位の典型だ。')

add_separator()

# ===== 第3章 =====
doc.add_heading('第3章：甘口の定着、そして辛口復権の兆し', level=1)

add_text('SAKETIMEのレビューデータは2017年〜2025年をカバーしている。この8年間の変化を追うと、単純な「辛口から甘口へ」では語れない、もう少し複雑な物語が見えてくる。')

doc.add_heading('構造化データの時系列推移', level=2)

add_image('fig6_trend_structured.png', '図7: テイスト傾向の時系列推移（構造化データ、全銘柄）')

add_text('4つの指標を同じスケールで並べると、興味深い動きが読み取れる。')

add_bullet('甘口率は40%前後で最も高い位置を維持しているが、実は2017年の42%から2025年の37%へ緩やかに低下している。甘口は王道として定着したものの、独走状態ではなくなりつつある。')
add_bullet('辛口率は2017年の22%から2022年の18%まで低下が続いた。しかし2023年を境に反転し、2025年には20%台へ再上昇している。辛口復権の兆しだ。')
add_bullet('濃醇率は20%→14%へ着実に低下を続けており、「重い酒」離れのトレンドには変化がない。')
add_bullet('軽快率は28-30%で安定しており、大きな変動はない。')

add_text('つまり、2017年から2022年にかけては確かに「甘口優位・辛口退潮」の流れがあった。しかし直近の2023年以降、辛口率が反転上昇し、甘口率がやや下がっている。甘口一辺倒の時代は、すでにピークを過ぎた可能性がある。')

doc.add_heading('テキスト分析の時系列推移', level=2)

add_image('fig7_trend_text.png', '図8: 味わいキーワード出現率の推移（テキスト分析、全銘柄）')

add_text('レビューテキストの分析でも同様の傾向が確認できる。最も顕著な変化はフルーティ系キーワードの増加で、2017年の17%から2025年には25%まで上昇した。フルーティという表現は現在進行形で定着が進んでいる。')

add_text('一方、甘口系キーワードは2019年頃に一段階上がったが、その後は横ばいだ。「いま強まっている」というよりは「既に主流として定着した」と言える。')

add_quote('甘口は王道として定着した。しかし辛口も復権しつつある。一方的なシフトではなく、多様化が進んでいるのかもしれない。')

add_separator()

# ===== 補章 =====
doc.add_heading('補章：別のデータソースで裏付ける ― さけのわのフレーバー分析', level=1)

add_text('ここまでの分析はすべてSAKETIMEのデータに基づいている。結論の妥当性を検証するため、独立したデータソースである「さけのわ」のデータでクロスチェックを行った。')

add_text('さけのわは日本酒コミュニティアプリで、100万件以上のチェックイン記録を持つ。公式に公開されているAPIから、約3,100銘柄のフレーバーデータを取得できる。特に有用なのは、レビューテキストのNLP解析から算出された「フレーバーチャート」（華やか/芳醇/重厚/穏やか/軽快/ドライの6軸数値）と、銘柄ごとに付与されたフレーバータグ（141種類）だ。')

add_note('【データの性質について】さけのわのフレーバータグはSAKETIMEのテイスト評価（甘い/辛いの排他的選択）とは性質が異なる。レビューテキストから抽出された特徴語の集積であり、1銘柄に平均9.4個のタグが付く。「甘味」と「辛口」が同時に付いている銘柄も30%存在する。したがって、タグの出現率はその味わいが言及される頻度を示すものであり、排他的な分類ではない。また、さけのわデータプロジェクト (https://sakenowa.com) のデータを加工して利用している。')

add_image('fig12_sakenowa_analysis.png', '図：さけのわランキング帯別のフレーバータグ出現率（左）とフレーバーチャート6軸平均（右）')

doc.add_heading('フレーバーチャート6軸が語るもの', level=2)

add_text('右のレーダーチャートは、ランキング帯別にフレーバーチャート6軸の平均値をプロットしたものだ。上位銘柄と圏外銘柄で形が明確に異なる。')

add_bullet('華やか: 上位ほど高い（1-10位: 0.50 → 圏外: 0.35）。上位銘柄の最大の特徴は「華やかさ」にある')
add_bullet('重厚: 上位ほど低い（1-10位: 0.24 → 圏外: 0.35）。SAKETIMEで見た「上位は芳醇ではなく軽め」という結論と完全に一致する')
add_bullet('ドライ: 意外にも上位ほどやや高い（1-10位: 0.47 → 圏外: 0.39）。これは「辛口」とは異なる文脈で、キレの良さや洗練された味わいを反映している可能性がある')

doc.add_heading('フレーバータグが語るもの', level=2)

add_text('左の棒グラフでは、主要なフレーバータグの出現率をランキング帯別に比較した。')

add_bullet('辛口タグ: 1-10位で0%、51-100位で78%。SAKETIMEの分析（Top10で辛口率4%、下位で29%）と方向性が完全に一致')
add_bullet('フルーティ: 1-10位で100%、圏外で33%。上位銘柄のフルーティ傾向がここでも裏付けられた')
add_bullet('濃厚・コク: 上位ほど低い。「芳醇」が上位の特徴ではないことを再確認')
add_bullet('バランス・フレッシュ・さわやか: 上位で突出して高い。上位銘柄の本質は「重さ」ではなく「調和とみずみずしさ」にある')

add_quote('独立したデータソースで、SAKETIMEの結論が裏付けられた。上位銘柄は「華やかでフルーティ、軽やかでバランスが良い」。辛口や濃厚は上位ほど少ない。')

add_separator()

# ===== 第4章 =====
doc.add_heading('第4章：では「辛口の時代」は終わったのか？ ― 意外な答え', level=1)

add_text('SAKETIMEとさけのわ、2つのデータソースから同じ傾向が確認された。では、もっと広い視点で見るとどうか。一般消費者はどう動いているのだろう。')

doc.add_heading('Google Trendsが示す一般消費者のリアル', level=2)

add_text('Google Trendsで2004年以降の検索トレンドを確認した。NDLのデータとは異なり、Google Trendsでは「日本酒」との組み合わせで検索しているため、日本酒の文脈に限定されたデータである。')

add_image('fig9_google_trends.png', '図9: Google Trendsにおける日本酒関連の検索トレンド（2004〜2025年）')

add_text('結果は明快だった。Google検索では「日本酒 辛口」が「日本酒 甘口」の約2.5倍の検索量を一貫して維持している。しかも、辛口の検索量は2004年から2025年にかけて右肩上がりで増え続けている。')

add_text('SAKETIMEの愛好家コミュニティでは甘口・フルーティが上位を占めていたが、一般消費者が日本酒を探すときの軸は、いまだに「辛口」なのだ。')

doc.add_heading('三つの層', level=2)

add_quote('辛口の時代は「終わっていない」。むしろ、日本酒の味わいをめぐる風景は三層構造になっている。')

add_bullet('第1層・ランキング上位（Top10〜50位）：甘口・フルーティ・軽快が支配的。甘口率67%、辛口率4%。')
add_bullet('第2層・ランキング下位（501位以降）：甘口率28%、辛口率29%。甘辛が拮抗しており、辛口の銘柄も数多く存在。')
add_bullet('第3層・一般消費者（Google検索）：「辛口」が一貫して主要キーワード。検索量は増加中。')

add_text('しかも、第3章で見たように、SAKETIMEのレビュー全体でも辛口率は近年再上昇している。愛好家コミュニティの中でさえ、辛口への揺り戻しが始まっているのだ。「辛口がいい酒」は終わるどころか、むしろしぶとく生き続けている。')

add_separator()

# ===== まとめ =====
doc.add_heading('まとめ：140年のサイクルと「辛口」のしぶとさ', level=1)

add_text('3つのデータソースを時系列で接続すると、日本酒の味わいトレンドの大きなサイクルが浮かび上がる。ただし、各データソースには性質の違い（NDLは全出版物の語頻度、Google Trendsは検索語の組み合わせ、SAKETIMEは愛好家のレビュー）があり、単純な比較には注意が必要だ。その点を踏まえた上で、大きな流れとしては以下のように読み取れる。')

add_bullet('1880年代〜：甘口の時代（出版物で「甘口」が圧倒的）')
add_bullet('1970年代〜90年代：淡麗辛口ブーム（「辛口」「吟醸」が急上昇）')
add_bullet('2000年代〜：純米吟醸ブーム（Google検索で急上昇）')
add_bullet('2010年代後半〜：甘口・フルーティが王道化（SAKETIMEレビューで顕著）')
add_bullet('2023年〜：辛口復権の兆し（SAKETIME全体で辛口率が再上昇）')

add_text('約140年をかけて、甘口→辛口→甘口（フルーティ）という大きな揺り戻しが起きてきた。しかし、現在起きていることは単純な「甘口回帰」ではない。甘口・フルーティが王道として定着する一方で、辛口も復権しつつあり、味わいの多様化が進んでいるように見える。')

add_text('そして最も印象的な発見は、「辛口がいい酒」のしぶとさだ。SAKETIMEのランキング上位では甘口が圧倒的でも、501位以降では辛口率が29%を占める。一般消費者のGoogle検索では辛口が甘口の2.5倍。そしてSAKETIMEのレビュー全体でも、辛口率は近年再び上昇に転じている。')

add_text('「辛口がいい酒」の時代は終わったか？ ― データが示す答えは、「終わっていない。そしてこれからも終わらないかもしれない」である。')

add_separator()

# ===== 分析手法 =====
doc.add_heading('分析手法について', level=1)
add_bullet('SAKETIMEデータ: Pythonによるスクレイピング（BeautifulSoup）。5,386銘柄のランキング情報、銘柄詳細、約37万件のレビューを取得')
add_bullet('テキスト分析: キーワード出現率の集計、ランキング帯別・年代別クロス分析')
add_bullet('さけのわデータ: さけのわデータプロジェクトAPI（約3,100銘柄のフレーバーチャート6軸・フレーバータグ141種）')

add_note('【レビュアーの偏りについて】SAKETIMEの約37万件のレビューは4,625人のユーザーによって投稿されているが、その分布は極めて偏っている。最多投稿者は1人で約14,000件（全体の3.8%）を投稿しており、上位100人（全ユーザーの2.2%）が全レビューの37%を、上位500人（10.8%）が72%を占める。本分析の結果は、こうした一部のヘビーレビュアーの嗜好に強く影響されている可能性がある点に留意されたい。')
add_bullet('出版物データ: 国立国会図書館 NDL Ngram Viewer API（図書・雑誌約230万点のOCRテキスト）')
add_bullet('検索トレンド: Google Trends（pytrends、地域:日本、2004年〜2025年）')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('データ取得日: 2026年4月14日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(130, 130, 130)
run.font.name = 'Yu Gothic'

# 保存
output_path = os.path.join(BLOG_DIR, 'article.docx')
doc.save(output_path)
print(f'保存完了: {output_path}')
